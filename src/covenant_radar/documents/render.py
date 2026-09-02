"""Render persisted warning memos into committee-pack artefacts.

The renderer is deliberately a presentation adapter.  It accepts a stored
memo-shaped object and an explicit export context; it does not calculate a
ratio, rewrite model prose, or infer a missing value.  The canonical
integrity digest is calculated from the memo's persisted content and is
therefore stable when the same memo is exported more than once.  The export
timestamp is metadata about the rendering event, not part of the memo's
meaning.
"""

from __future__ import annotations

import hashlib
import json
import os
from collections.abc import Mapping, Sequence
from dataclasses import dataclass, field
from datetime import UTC, date, datetime
from decimal import Decimal
from io import BytesIO
from math import isfinite
from pathlib import Path
from typing import Any, Final, TypedDict
from uuid import UUID
from zoneinfo import ZoneInfo

from jinja2 import Environment, FileSystemLoader, StrictUndefined, select_autoescape

_TEMPLATE_ROOT: Final[Path] = Path(__file__).resolve().parents[1] / "web" / "templates"
_WEB_ROOT: Final[Path] = _TEMPLATE_ROOT.parent
_IST: Final[ZoneInfo] = ZoneInfo("Asia/Kolkata")
_SHA256_LENGTH: Final[int] = hashlib.sha256().digest_size * 2
_MAX_RENDERED_TEXT_BYTES: Final[int] = 2 * 1024 * 1024
_MAX_COLLECTION_ITEMS: Final[int] = 500
_MAX_NESTING: Final[int] = 20
_MODEL_DRAFT_LABEL: Final[str] = "Model-drafted connecting prose"
_SLOT_ORDER: Final[tuple[str, ...]] = (
    "situation",
    "ratio_name",
    "value",
    "threshold",
    "headroom",
    "probability",
    "confidence",
    "crossing_date",
    "drivers",
    "evidence_counts",
    "simulation_options",
    "recommended_interventions",
    "intervention_text",
)
_SLOT_LABELS: Final[Mapping[str, str]] = {
    "situation": "Situation",
    "ratio_name": "Covenant",
    "value": "Recorded value",
    "threshold": "Threshold in force",
    "headroom": "Headroom",
    "probability": "Projected breach probability",
    "confidence": "Forecast confidence",
    "crossing_date": "Projected crossing date",
    "drivers": "Drivers",
    "evidence_counts": "Evidence citations",
    "simulation_options": "Simulation options",
    "recommended_interventions": "Recommended interventions",
    "intervention_text": "Intervention wording",
}
_DEFAULT_LETTERHEAD_NAME: Final[str] = "Covenant Radar"


@dataclass(frozen=True, slots=True)
class MemoLetterhead:
    """Validated letterhead settings used by both output formats."""

    name: str = _DEFAULT_LETTERHEAD_NAME
    address: str | None = None
    subtitle: str | None = None

    def __post_init__(self) -> None:
        object.__setattr__(self, "name", _required_text(self.name, "letterhead name", 200))
        if self.address is not None:
            object.__setattr__(
                self,
                "address",
                _optional_text(self.address, "letterhead address", 500),
            )
        if self.subtitle is not None:
            object.__setattr__(
                self,
                "subtitle",
                _optional_text(self.subtitle, "letterhead subtitle", 200),
            )

    @classmethod
    def from_value(cls, value: MemoLetterhead | Mapping[str, object] | None) -> MemoLetterhead:
        """Accept configuration mappings without allowing arbitrary fields."""

        if value is None:
            return cls()
        if isinstance(value, cls):
            return value
        if not isinstance(value, Mapping):
            raise TypeError("letterhead must be a MemoLetterhead, mapping, or None.")
        allowed = {"name", "organisation", "organization", "address", "subtitle"}
        unknown = set(value).difference(allowed)
        if unknown:
            raise ValueError(f"Unknown letterhead setting: {sorted(unknown)[0]!r}.")
        names = [value[key] for key in ("name", "organisation", "organization") if key in value]
        if len(names) > 1 and len({str(item) for item in names}) != 1:
            raise ValueError("letterhead name settings disagree.")
        name = names[0] if names else _DEFAULT_LETTERHEAD_NAME
        if not isinstance(name, str):
            raise TypeError("letterhead name must be text.")
        address = value.get("address")
        if address is not None and not isinstance(address, str):
            raise TypeError("letterhead address must be text or None.")
        subtitle = value.get("subtitle")
        if subtitle is not None and not isinstance(subtitle, str):
            raise TypeError("letterhead subtitle must be text or None.")
        return cls(
            name=name,
            address=address,
            subtitle=subtitle,
        )


@dataclass(frozen=True, slots=True)
class MemoExportContext:
    """Per-render metadata that is intentionally separate from memo content."""

    memo_id: UUID | str
    integrity_hash: str
    exported_at: datetime
    exported_by: str
    letterhead: MemoLetterhead = field(default_factory=MemoLetterhead)

    def __post_init__(self) -> None:
        if not isinstance(self.memo_id, UUID | str) or not str(self.memo_id).strip():
            raise TypeError("MemoExportContext.memo_id must be a non-empty UUID or string.")
        digest = self.integrity_hash.lower() if isinstance(self.integrity_hash, str) else ""
        if len(digest) != _SHA256_LENGTH or any(
            character not in "0123456789abcdef" for character in digest
        ):
            raise ValueError("MemoExportContext.integrity_hash must be a SHA-256 hex digest.")
        if not isinstance(self.exported_at, datetime):
            raise TypeError("MemoExportContext.exported_at must be a datetime.")
        if self.exported_at.tzinfo is None or self.exported_at.utcoffset() is None:
            raise ValueError("MemoExportContext.exported_at must be timezone-aware.")
        object.__setattr__(self, "integrity_hash", digest)
        object.__setattr__(self, "exported_at", self.exported_at.astimezone(UTC))
        object.__setattr__(
            self,
            "exported_by",
            _required_text(self.exported_by, "exported_by", 200),
        )
        object.__setattr__(self, "letterhead", MemoLetterhead.from_value(self.letterhead))

    @property
    def exported_at_ist(self) -> str:
        """Return the export instant in the required Indian time zone."""

        return self.exported_at.astimezone(_IST).strftime("%d %b %Y, %H:%M:%S IST")


@dataclass(frozen=True, slots=True)
class RenderedField:
    """A safe, already-stringified field for the HTML and DOCX adapters."""

    label: str
    value: str


class RenderedSlot(TypedDict):
    """A string-only slot prepared for both export adapters."""

    name: str
    label: str
    value: str
    state: str
    sources: str


@dataclass(frozen=True, slots=True)
class RenderedCollection:
    """A collection of stored mapping fields, retaining every input field."""

    title: str
    items: tuple[tuple[RenderedField, ...], ...]


class MemoDocumentData(TypedDict):
    """Typed context shared by the HTML template and DOCX adapter."""

    letterhead: MemoLetterhead
    metadata: tuple[tuple[str, str], ...]
    integrity_hash: str
    exported_at_ist: str
    exported_by: str
    slots: tuple[RenderedSlot, ...]
    draft_paragraphs: tuple[str, ...]
    collections: tuple[RenderedCollection, ...]


class MemoRenderer:
    """Render one stored memo through the fixed export template."""

    def __init__(self, *, template_directory: Path | str = _TEMPLATE_ROOT) -> None:
        directory = Path(template_directory).expanduser().resolve()
        if not directory.is_dir():
            raise FileNotFoundError(f"Memo export template directory does not exist: {directory}")
        template_path = directory / "exports" / "memo.html"
        if not template_path.is_file():
            raise FileNotFoundError(f"Memo export template does not exist: {template_path}")
        self.template_directory = directory
        self.web_root = directory.parent
        self.environment = Environment(
            loader=FileSystemLoader(str(directory)),
            autoescape=select_autoescape(("html", "xml")),
            undefined=StrictUndefined,
        )

    def render_html(self, memo: object, context: MemoExportContext) -> str:
        """Render the canonical HTML source used by PDF and browser previews."""

        if not isinstance(context, MemoExportContext):
            raise TypeError("MemoRenderer.render_html requires a MemoExportContext.")
        document = _document_data(memo, context)
        html = self.environment.get_template("exports/memo.html").render(**document)
        if len(html.encode("utf-8")) > _MAX_RENDERED_TEXT_BYTES:
            raise ValueError("Memo export exceeds the maximum rendered document size.")
        return html

    def render_pdf(self, memo: object, context: MemoExportContext) -> bytes:
        """Render a paginated PDF with a repeated provenance footer."""

        html_type, dll_handles = _weasyprint_html_type()
        try:
            html = self.render_html(memo, context)
            content = html_type(string=html, base_url=str(self.web_root)).write_pdf()
        finally:
            for handle in dll_handles:
                handle.close()
        if not isinstance(content, bytes) or not content:
            raise RuntimeError("PDF memo export produced no content.")
        return content

    def render_docx(self, memo: object, context: MemoExportContext) -> bytes:
        """Render a DOCX containing the same stored fields and assumptions."""

        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches, Pt
        except ImportError as error:
            raise RuntimeError(
                "DOCX memo export requires the pinned python-docx dependency."
            ) from error

        document = _document_data(memo, context)
        output = Document()
        section = output.sections[0]
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

        normal = output.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(10)

        _add_docx_header(section, document["letterhead"], WD_ALIGN_PARAGRAPH, Pt)
        _add_docx_footer(section, WD_ALIGN_PARAGRAPH)

        title = output.add_heading("Covenant warning memo", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _add_docx_metadata(output, document["metadata"])

        output.add_heading("Recorded figures and provenance", level=1)
        slots_table = output.add_table(rows=1, cols=4)
        slots_table.style = "Table Grid"
        _set_row(slots_table.rows[0], ("Figure", "Value", "State", "Source record(s)"))
        for slot in document["slots"]:
            row = slots_table.add_row()
            _set_row(row, (slot["label"], slot["value"], slot["state"], slot["sources"]))

        output.add_heading(_MODEL_DRAFT_LABEL, level=1)
        for paragraph in document["draft_paragraphs"]:
            output.add_paragraph(paragraph)

        _add_docx_collections(output, document["collections"])
        output.add_paragraph(
            "This memo is advisory. Human credit review is required before action."
        )

        properties = output.core_properties
        properties.title = "Covenant warning memo"
        properties.subject = f"Integrity hash {context.integrity_hash}"
        properties.author = context.exported_by
        created_at = _memo_datetime(memo, "created_at") or context.exported_at
        properties.created = created_at
        properties.modified = created_at

        stream = BytesIO()
        output.save(stream)
        content = stream.getvalue()
        if not content:
            raise RuntimeError("DOCX memo export produced no content.")
        return content

    def render(self, memo: object, context: MemoExportContext, format: str) -> bytes:
        """Render ``format`` after strict format validation."""

        normalized = _format(format)
        if normalized == "pdf":
            return self.render_pdf(memo, context)
        return self.render_docx(memo, context)


def memo_integrity_hash(memo: object) -> str:
    """Hash all persisted memo content that an export presents.

    Export timestamps, exporter identity, and the output container are not
    included.  Consequently the digest identifies the memo's meaning rather
    than one rendering event, while the database export row identifies every
    rendering event separately.
    """

    payload = {
        "memo_id": _json_value(_memo_attr(memo, "id")),
        "borrower_id": _json_value(_memo_attr(memo, "borrower_id")),
        "run_id": _json_value(_memo_attr(memo, "run_id")),
        "case_id": _json_value(_memo_attr(memo, "case_id")),
        "template_version": _json_value(_memo_attr(memo, "template_version")),
        "prompt_version": _json_value(_memo_attr(memo, "prompt_version")),
        "provider": _json_value(_memo_attr(memo, "provider")),
        "model_version": _json_value(_memo_attr(memo, "model_version")),
        "slots": _json_value(_memo_attr(memo, "slots")),
        "drafted_text": _json_value(_memo_attr(memo, "drafted_text")),
        "actions": _json_value(_memo_attr(memo, "actions")),
        "simulations": _json_value(_memo_attr(memo, "simulations")),
        "check_verdict": _json_value(_memo_attr(memo, "check_verdict")),
        "generated_by_id": _json_value(_memo_attr(memo, "generated_by_id")),
        "created_at": _json_value(_memo_attr(memo, "created_at")),
    }
    encoded = json.dumps(
        payload,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
        allow_nan=False,
    ).encode("utf-8")
    return hashlib.sha256(encoded).hexdigest()


def render_memo_pdf(
    memo: object,
    context: MemoExportContext,
    *,
    renderer: MemoRenderer | None = None,
) -> bytes:
    """Functional PDF rendering entry point."""

    return (renderer or MemoRenderer()).render_pdf(memo, context)


def render_memo_docx(
    memo: object,
    context: MemoExportContext,
    *,
    renderer: MemoRenderer | None = None,
) -> bytes:
    """Functional DOCX rendering entry point."""

    return (renderer or MemoRenderer()).render_docx(memo, context)


def _document_data(memo: object, context: MemoExportContext) -> MemoDocumentData:
    slots = _render_slots(_memo_attr(memo, "slots"))
    simulations = _render_collection(_memo_attr(memo, "simulations"), "Simulations")
    actions = _render_collection(_memo_attr(memo, "actions"), "Actions in memo")
    collections = tuple(item for item in (simulations, actions) if item.items)
    drafted_text = _required_text(_memo_attr(memo, "drafted_text"), "memo drafted_text", 200_000)
    paragraphs = tuple(part.strip() for part in drafted_text.split("\n\n") if part.strip())
    metadata = (
        ("Memo id", str(context.memo_id)),
        ("Generated at (IST)", context.exported_at_ist),
        ("Generated by", context.exported_by),
        ("Integrity hash (SHA-256)", context.integrity_hash),
        ("Template version", _text_or_absent(_memo_attr(memo, "template_version"))),
        ("Model-drafted sections", _MODEL_DRAFT_LABEL),
    )
    return {
        "letterhead": context.letterhead,
        "metadata": metadata,
        "integrity_hash": context.integrity_hash,
        "exported_at_ist": context.exported_at_ist,
        "exported_by": context.exported_by,
        "slots": slots,
        "draft_paragraphs": paragraphs,
        "collections": collections,
    }


def _render_slots(raw_slots: object) -> tuple[RenderedSlot, ...]:
    if not isinstance(raw_slots, Mapping):
        raise ValueError("Memo slots must be a mapping.")
    nested = raw_slots.get("slots", raw_slots)
    if not isinstance(nested, Mapping):
        raise ValueError("Memo slots.slots must be a mapping.")
    if len(nested) > _MAX_COLLECTION_ITEMS:
        raise ValueError("Memo slots contain too many entries.")
    entries: list[tuple[str, object]] = []
    for name in _SLOT_ORDER:
        if name in nested:
            entries.append((name, nested[name]))
    entries.extend((str(name), value) for name, value in nested.items() if name not in _SLOT_ORDER)
    if not entries:
        raise ValueError("Memo slots must contain at least one slot.")
    rendered: list[RenderedSlot] = []
    for name, raw_slot in entries:
        if not isinstance(raw_slot, Mapping):
            raise ValueError(f"Memo slot {name!r} must be a mapping.")
        state = _text_or_absent(raw_slot.get("state"))
        value = _display_value(raw_slot.get("value"), field=f"slot {name}")
        reason = raw_slot.get("reason")
        if reason is not None:
            value = f"{value} ({_display_value(reason, field=f'slot {name} reason')})"
        references = raw_slot.get("record_references", ())
        source_text = _references_text(references)
        rendered.append(
            RenderedSlot(
                name=name,
                label=_SLOT_LABELS.get(name, _humanize(name)),
                value=value,
                state=state,
                sources=source_text,
            )
        )
    return tuple(rendered)


def _render_collection(raw_collection: object, title: str) -> RenderedCollection:
    if raw_collection is None:
        return RenderedCollection(title, ())
    raw_items: object = raw_collection
    if isinstance(raw_collection, Mapping):
        raw_items = raw_collection.get("items", ())
    if isinstance(raw_items, str | bytes | bytearray) or not isinstance(raw_items, Sequence):
        raise ValueError(f"Memo {title.lower()} must contain a sequence of items.")
    if len(raw_items) > _MAX_COLLECTION_ITEMS:
        raise ValueError(f"Memo {title.lower()} contains too many items.")
    items: list[tuple[RenderedField, ...]] = []
    for index, raw_item in enumerate(raw_items):
        if isinstance(raw_item, Mapping):
            rendered_fields: list[RenderedField] = []
            for key, value in raw_item.items():
                field_name = str(key)
                if (
                    field_name == "assumptions"
                    and isinstance(value, Sequence)
                    and not isinstance(value, str | bytes | bytearray)
                ):
                    if len(value) > _MAX_COLLECTION_ITEMS:
                        raise ValueError(
                            f"Memo {title.lower()} assumptions contain too many items."
                        )
                    rendered_value = "\n".join(
                        _display_value(item, field=f"{title}[{index}].assumptions")
                        for item in value
                    )
                else:
                    rendered_value = _display_value(value, field=f"{title}[{index}]")
                rendered_fields.append(RenderedField(_humanize(field_name), rendered_value))
            fields = tuple(rendered_fields)
        else:
            fields = (RenderedField("Value", _display_value(raw_item, field=f"{title}[{index}]")),)
        items.append(fields)
    return RenderedCollection(title, tuple(items))


def _references_text(raw_references: object) -> str:
    if raw_references is None:
        return "Not available"
    if isinstance(raw_references, Mapping):
        references: Sequence[object] = (raw_references,)
    elif isinstance(raw_references, str | bytes | bytearray):
        references = (raw_references,)
    elif isinstance(raw_references, Sequence):
        references = raw_references
    else:
        raise ValueError("Memo slot record_references must be a sequence.")
    if not references:
        return "Not available"
    if len(references) > _MAX_COLLECTION_ITEMS:
        raise ValueError("Memo slot record_references contains too many entries.")
    result: list[str] = []
    for reference in references:
        if isinstance(reference, Mapping):
            record_type = _text_or_absent(reference.get("type", reference.get("record_type")))
            record_id = _text_or_absent(reference.get("id", reference.get("record_id")))
            result.append(f"{record_type}: {record_id}")
        else:
            result.append(_display_value(reference, field="record reference"))
    return ", ".join(result)


def _display_value(value: object, *, field: str) -> str:
    if isinstance(value, str):
        return value
    if isinstance(value, date | datetime | Decimal | UUID):
        return str(_json_value(value))
    if isinstance(value, bool | int | float) or value is None:
        return _text_or_absent(value)
    normalized = _json_value(value)
    return json.dumps(normalized, ensure_ascii=False, sort_keys=True, separators=(", ", ": "))


def _json_value(value: object, *, depth: int = 0) -> Any:
    if depth > _MAX_NESTING:
        raise ValueError("Memo export data is nested too deeply.")
    if value is None or isinstance(value, str | bool | int):
        return value
    if isinstance(value, float):
        if not isfinite(value):
            raise ValueError("Memo export data contains a non-finite float.")
        return value
    if isinstance(value, Decimal | UUID):
        return str(value)
    if isinstance(value, datetime | date):
        return value.isoformat()
    if isinstance(value, Mapping):
        if len(value) > _MAX_COLLECTION_ITEMS:
            raise ValueError("Memo export mapping contains too many entries.")
        normalized: dict[str, Any] = {}
        for key, item in value.items():
            if not isinstance(key, str):
                raise TypeError("Memo export mapping keys must be strings.")
            normalized[key] = _json_value(item, depth=depth + 1)
        return normalized
    if isinstance(value, Sequence) and not isinstance(value, str | bytes | bytearray):
        if len(value) > _MAX_COLLECTION_ITEMS:
            raise ValueError("Memo export sequence contains too many entries.")
        return [_json_value(item, depth=depth + 1) for item in value]
    raise TypeError(f"Memo export data contains unsupported {type(value).__name__}.")


def _memo_attr(memo: object, name: str) -> object:
    try:
        value = getattr(memo, name)
    except AttributeError as error:
        raise ValueError(f"Stored memo is missing {name}.") from error
    return value


def _memo_datetime(memo: object, name: str) -> datetime | None:
    value = _memo_attr(memo, name)
    if value is None:
        return None
    if not isinstance(value, datetime) or value.tzinfo is None or value.utcoffset() is None:
        raise ValueError(f"Stored memo {name} must be timezone-aware.")
    return value.astimezone(UTC)


def _format(value: object) -> str:
    if not isinstance(value, str):
        raise TypeError("Memo export format must be text.")
    normalized = value.strip().lower()
    if normalized not in {"pdf", "docx"}:
        raise ValueError(f"Unsupported memo export format: {value!r}.")
    return normalized


def _text_or_absent(value: object) -> str:
    if value is None:
        return "Not available from the recorded evidence."
    if isinstance(value, bool):
        return "true" if value else "false"
    if isinstance(value, float) and not isfinite(value):
        raise ValueError("Memo export text cannot contain a non-finite float.")
    return str(value)


def _required_text(value: object, field: str, maximum: int) -> str:
    if not isinstance(value, str) or not value.strip():
        raise ValueError(f"{field} must be non-blank text.")
    cleaned = value.strip()
    if len(cleaned) > maximum:
        raise ValueError(f"{field} must be at most {maximum} characters.")
    if any(ord(character) < 32 and character not in "\n\r\t" for character in cleaned):
        raise ValueError(f"{field} contains a control character.")
    return cleaned


def _optional_text(value: object, field: str, maximum: int) -> str | None:
    if value is None:
        return None
    return _required_text(value, field, maximum)


def _humanize(value: str) -> str:
    return value.replace("_", " ").strip().title()


def _weasyprint_html_type() -> tuple[Any, list[Any]]:
    """Load WeasyPrint and, on Windows, expose DLL directories from PATH."""

    handles = _windows_dll_handles() if os.name == "nt" else []
    try:
        from weasyprint import HTML  # type: ignore[import-untyped]

        return HTML, handles
    except (ImportError, OSError) as error:
        for handle in handles:
            handle.close()
        message = "PDF memo export requires the pinned weasyprint dependency."
        if os.name == "nt":
            message = (
                "PDF memo export requires the pinned weasyprint dependency and native libraries."
            )
        raise RuntimeError(message) from error


def _windows_dll_handles() -> list[Any]:
    """Register valid PATH directories for Python's Windows DLL loader."""

    handles: list[Any] = []
    for raw_directory in os.environ.get("PATH", "").split(os.pathsep):
        if not raw_directory:
            continue
        directory = Path(raw_directory)
        if not directory.is_dir():
            continue
        try:
            handles.append(os.add_dll_directory(str(directory)))
        except OSError:
            continue
    return handles


def _add_docx_header(section: Any, letterhead: MemoLetterhead, align: Any, point: Any) -> None:
    paragraph = section.header.paragraphs[0]
    paragraph.alignment = align.LEFT
    run = paragraph.add_run(letterhead.name)
    run.bold = True
    run.font.size = point(11)
    if letterhead.subtitle:
        paragraph.add_run(f"\n{letterhead.subtitle}")
    if letterhead.address:
        paragraph.add_run(f"\n{letterhead.address}")


def _add_docx_footer(section: Any, align: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = align.CENTER
    paragraph.add_run(f"{_MODEL_DRAFT_LABEL}; figures are record-backed.  Page ")
    page_field = OxmlElement("w:fldSimple")
    page_field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(page_field)
    paragraph.add_run(" of ")
    pages_field = OxmlElement("w:fldSimple")
    pages_field.set(qn("w:instr"), "NUMPAGES")
    paragraph._p.append(pages_field)


def _add_docx_metadata(document: Any, metadata: Sequence[tuple[str, str]]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in metadata:
        row = table.add_row()
        _set_row(row, (label, value))


def _add_docx_collections(document: Any, collections: Sequence[RenderedCollection]) -> None:
    for collection in collections:
        document.add_heading(collection.title, level=1)
        for index, item in enumerate(collection.items, start=1):
            item_title = (
                collection.title[:-1] if collection.title.endswith("s") else collection.title
            )
            document.add_heading(f"{item_title} {index}", level=2)
            table = document.add_table(rows=0, cols=2)
            table.style = "Table Grid"
            for rendered_field in item:
                row = table.add_row()
                _set_row(row, (rendered_field.label, rendered_field.value))


def _set_row(row: Any, values: Sequence[str]) -> None:
    if len(row.cells) != len(values):
        raise ValueError("DOCX row and value counts do not match.")
    for cell, value in zip(row.cells, values, strict=True):
        cell.text = value


# Public aliases make the adapter discoverable to callers using either noun.
MemoExportRenderer = MemoRenderer
MemoLetterheadConfig = MemoLetterhead
ExportContext = MemoExportContext
integrity_hash_for_memo = memo_integrity_hash


__all__ = [
    "ExportContext",
    "MemoExportContext",
    "MemoExportRenderer",
    "MemoLetterhead",
    "MemoLetterheadConfig",
    "MemoRenderer",
    "RenderedCollection",
    "RenderedField",
    "integrity_hash_for_memo",
    "memo_integrity_hash",
    "render_memo_docx",
    "render_memo_pdf",
]

"""Integration coverage for T-102 memo PDF and DOCX exports."""

from __future__ import annotations

import hashlib
from collections.abc import Mapping
from datetime import UTC, datetime, timedelta
from io import BytesIO
from uuid import uuid4

import pytest
from docx import Document
from pypdf import PdfReader
from sqlalchemy import create_engine, func, select
from sqlalchemy.orm import Session

from covenant_radar.core.clock import FixedClock
from covenant_radar.core.errors import AuthorizationError
from covenant_radar.db.base import Base
from covenant_radar.db.models.borrower import Borrower
from covenant_radar.db.models.identity import AppUser
from covenant_radar.db.models.portfolio import Portfolio
from covenant_radar.db.models.workflow import Memo, MemoExport
from covenant_radar.db.scoping import Scope
from covenant_radar.security.permissions import Permission
from covenant_radar.security.rbac import Principal
from covenant_radar.services.memo import MemoExportService

pytestmark = pytest.mark.integration

_NOW = datetime(2026, 8, 31, 9, 30, tzinfo=UTC)


class _ExportStore:
    def __init__(self) -> None:
        self.objects: dict[str, bytes] = {}
        self.put_calls = 0

    def put(self, content: bytes) -> str:
        self.put_calls += 1
        key = f"memo-exports/sha256/{hashlib.sha256(content).hexdigest()}"
        self.objects[key] = content
        return key

    def delete(self, storage_key: str) -> None:
        del self.objects[storage_key]


class _Fixture:
    def __init__(self) -> None:
        self.engine = create_engine("sqlite:///:memory:")
        Base.metadata.create_all(self.engine)
        self.session = Session(self.engine)
        self.store = _ExportStore()
        self.user = AppUser(
            id=uuid4(),
            username="memo-exporter",
            email="memo-exporter@example.test",
            full_name="Memo Exporter",
            auth_source="local",
            is_active=True,
            created_at=_NOW,
            updated_at=_NOW,
            request_id="rq-t102-user",
        )
        self.portfolio = Portfolio.create(
            code="T102",
            name="Memo export portfolio",
            created_at=_NOW,
            updated_at=_NOW,
            request_id="rq-t102-portfolio",
        )
        self.borrower = Borrower(
            id=uuid4(),
            reference="B-T102-001",
            legal_name="T102 Borrower Private Limited",
            portfolio_id=self.portfolio.id,
            created_at=_NOW,
            updated_at=_NOW,
            request_id="rq-t102-borrower",
        )
        self.session.add_all([self.user, self.portfolio, self.borrower])
        self.session.flush()
        self.memo = Memo(
            id=uuid4(),
            borrower_id=self.borrower.id,
            template_version="v1",
            prompt_version="v2",
            provider="fixture",
            model_version="fixture-model",
            slots=_slots(),
            drafted_text=(
                "Debt service coverage is projected to reach the action point on 2026-10-15."
                "\n\nThe recorded value is 1.25 against a threshold of 1.10, with headroom of "
                "0.15. The projected breach probability is 0.42 at confidence 0.88."
                "\n\nReview and reduce funded exposure."
                "\n\nhuman credit review is required before action"
            ),
            actions={"items": [{"id": "REDUCE_DRAWING", "role_tag": "credit"}]},
            simulations={
                "items": [
                    {
                        "code": "REDUCE_DRAWING",
                        "projected_cross_date": "2026-12-01",
                        "probability": "0.20",
                        "delta_days": 47,
                        "delta_probability": "-0.22",
                        "assumptions": [
                            "The approved limit reduction takes effect immediately.",
                            "No new drawdown is assumed during the simulated period.",
                        ],
                    }
                ]
            },
            check_verdict="passed",
            generated_by_id=self.user.id,
            created_at=_NOW,
            updated_at=_NOW,
            created_by_id=self.user.id,
            updated_by_id=self.user.id,
            request_id="rq-t102-memo",
        )
        self.session.add(self.memo)
        self.session.flush()
        self.principal = Principal.user(self.user.id, (Permission.GENERATE_MEMO,))
        self.scope = Scope.from_paths(self.principal.id, [self.portfolio.path])

    def service(self, *, letterhead: Mapping[str, object] | None = None) -> MemoExportService:
        return MemoExportService(
            self.session,
            storage=self.store,
            clock=FixedClock(_NOW),
            request_id="rq-t102-export",
            letterhead=letterhead,
        )

    def close(self) -> None:
        self.session.close()
        self.engine.dispose()


def _slots() -> dict[str, object]:
    def slot(value: object, *references: tuple[str, str]) -> dict[str, object]:
        return {
            "value": value,
            "state": "present",
            "reason": None,
            "record_references": [
                {"type": record_type, "id": record_id} for record_type, record_id in references
            ],
        }

    return {
        "template_version": "v1",
        "slots": {
            "situation": slot("Projected pressure requires review.", ("triage", "triage-1")),
            "ratio_name": slot("Debt service coverage", ("forecast", "forecast-1")),
            "value": slot("1.25", ("forecast", "forecast-1")),
            "threshold": slot("1.10", ("forecast", "forecast-1")),
            "headroom": slot("0.15", ("forecast", "forecast-1")),
            "probability": slot("0.42", ("forecast", "forecast-1")),
            "confidence": slot("0.88", ("forecast", "forecast-1")),
            "crossing_date": slot("2026-10-15", ("forecast", "forecast-1")),
            "drivers": slot(
                [{"name": "Cash-flow pressure", "share": "0.60"}],
                ("forecast_driver", "driver-1"),
            ),
            "evidence_counts": slot(
                [{"citation": "EV-001", "count": 3}],
                ("evidence_item", "evidence-1"),
            ),
            "simulation_options": slot(
                [{"code": "REDUCE_DRAWING", "assumptions": ["full assumption"]}],
                ("simulation", "simulation-1"),
            ),
            "recommended_interventions": slot(
                [{"code": "REDUCE_DRAWING", "role_tag": "credit"}],
                ("intervention", "intervention-1"),
            ),
            "intervention_text": slot(
                ["Review and reduce funded exposure."],
                ("intervention", "intervention-1"),
            ),
        },
    }


def _pdf_text(content: bytes) -> str:
    return "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(content)).pages)


def _docx_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    values = [paragraph.text for paragraph in document.paragraphs]
    values.extend(
        paragraph.text for paragraph in document.sections[0].header.paragraphs if paragraph.text
    )
    values.extend(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    return "\n".join(values)


def _scalar_texts(value: object) -> tuple[str, ...]:
    if value is None:
        return ()
    if isinstance(value, Mapping):
        return tuple(text for item in value.values() for text in _scalar_texts(item))
    if isinstance(value, list | tuple):
        return tuple(text for item in value for text in _scalar_texts(item))
    return (str(value).lower() if isinstance(value, bool) else str(value),)


def _export(fixture: _Fixture, format: str = "pdf"):
    return fixture.service().export(
        fixture.memo,
        format=format,
        principal=fixture.principal,
        scope=fixture.scope,
    )


def test_pdf_and_docx_contain_same_figures_as_slots() -> None:
    fixture = _Fixture()
    try:
        pdf = _export(fixture, "pdf")
        docx = _export(fixture, "docx")
        pdf_text = _pdf_text(pdf.content)
        docx_text = _docx_text(docx.content)
        expected = tuple(
            text
            for name, value in _slots()["slots"].items()
            if name != "simulation_options"
            for text in _scalar_texts(value["value"])
        )
        for figure in expected:
            assert figure in pdf_text
            assert figure in docx_text
    finally:
        fixture.close()


def test_integrity_hash_stable_across_exports() -> None:
    fixture = _Fixture()
    try:
        first = _export(fixture, "pdf")
        later = _NOW + timedelta(minutes=2)
        second = fixture.service().export(
            fixture.memo,
            format="pdf",
            principal=fixture.principal,
            scope=fixture.scope,
            exported_at=later,
        )
        assert first.integrity_hash == second.integrity_hash
        assert first.record.integrity_hash == second.record.integrity_hash
        assert first.record.exported_at != second.record.exported_at
        assert first.content != second.content
    finally:
        fixture.close()


def test_assumptions_printed_in_full() -> None:
    fixture = _Fixture()
    try:
        pdf = _export(fixture, "pdf")
        docx = _export(fixture, "docx")
        for assumption in fixture.memo.simulations["items"][0]["assumptions"]:
            assert assumption in _pdf_text(pdf.content)
            assert assumption in _docx_text(docx.content)
    finally:
        fixture.close()


def test_default_letterhead_when_unconfigured() -> None:
    fixture = _Fixture()
    try:
        result = _export(fixture, "docx")
        assert "Covenant Radar" in _docx_text(result.content)
    finally:
        fixture.close()


def test_permission_enforced() -> None:
    fixture = _Fixture()
    try:
        unauthorized = Principal.user(fixture.user.id, ())
        with pytest.raises(AuthorizationError, match="GENERATE_MEMO"):
            fixture.service().export(
                fixture.memo,
                format="pdf",
                principal=unauthorized,
                scope=fixture.scope,
            )
        assert fixture.store.put_calls == 0
        assert fixture.session.scalar(select(func.count(MemoExport.id))) == 0
    finally:
        fixture.close()


def test_export_recorded() -> None:
    fixture = _Fixture()
    try:
        result = _export(fixture, "docx")
        row = fixture.session.scalar(select(MemoExport).where(MemoExport.id == result.record.id))
        assert row is not None
        assert row.memo_id == fixture.memo.id
        assert row.format == "docx"
        assert row.storage_key == result.storage_key
        assert row.integrity_hash == result.integrity_hash
        assert row.exported_by_id == fixture.user.id
    finally:
        fixture.close()
